"""
File Processor for extracting text from various file formats
"""

import logging
from typing import Optional
from pathlib import Path
import asyncio

# Import document processing libraries
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

logger = logging.getLogger(__name__)


class FileProcessor:
    """Process files and extract text content"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._extract_text_from_txt,
            '.md': self._extract_text_from_txt,
            '.pdf': self._extract_text_from_pdf,
            '.docx': self._extract_text_from_docx
        }
    
    async def extract_text(self, file_path: str) -> Optional[str]:
        """Extract text content from a file"""
        try:
            path = Path(file_path)
            if not path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            file_extension = path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                logger.warning(f"Unsupported file format: {file_extension}")
                return None
            
            # Get the appropriate extraction method
            extract_method = self.supported_formats[file_extension]
            
            # Run extraction in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            text_content = await loop.run_in_executor(None, extract_method, file_path)
            
            if text_content:
                # Clean up the text
                text_content = self._clean_text(text_content)
                logger.info(f"Extracted {len(text_content)} characters from {path.name}")
                return text_content
            else:
                logger.warning(f"No text content extracted from {path.name}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    async def extract_text_from_memory(self, file_content: bytes, file_name: str, mime_type: str) -> Optional[str]:
        """Extract text content from file bytes in memory"""
        try:
            file_extension = Path(file_name).suffix.lower()
            
            if file_extension not in self.supported_formats:
                logger.warning(f"Unsupported file format: {file_extension}")
                return None
            
            # Run extraction in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            
            if file_extension in ['.txt', '.md']:
                text_content = await loop.run_in_executor(None, self._extract_text_from_bytes_txt, file_content)
            elif file_extension == '.pdf':
                text_content = await loop.run_in_executor(None, self._extract_text_from_bytes_pdf, file_content)
            elif file_extension == '.docx':
                text_content = await loop.run_in_executor(None, self._extract_text_from_bytes_docx, file_content)
            else:
                return None
            
            if text_content:
                # Clean up the text
                text_content = self._clean_text(text_content)
                logger.info(f"Extracted {len(text_content)} characters from {file_name} (memory)")
                return text_content
            else:
                logger.warning(f"No text content extracted from {file_name}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_name} (memory): {str(e)}")
            return None
    
    def _extract_text_from_txt(self, file_path: str) -> Optional[str]:
        """Extract text from .txt or .md files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {str(e)}")
                return None
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return None
    
    def _extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF files"""
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 not available for PDF processing")
            return None
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return None
    
    def _extract_text_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX processing")
            return None
        
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return None
    
    def _extract_text_from_bytes_txt(self, file_content: bytes) -> Optional[str]:
        """Extract text from bytes (txt/md files)"""
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Try latin-1 as fallback
            try:
                return file_content.decode('latin-1')
            except Exception as e:
                logger.error(f"Error decoding text content: {str(e)}")
                return None
    
    def _extract_text_from_bytes_pdf(self, file_content: bytes) -> Optional[str]:
        """Extract text from PDF bytes"""
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 not available for PDF processing")
            return None
        
        try:
            import io
            text_content = []
            
            with io.BytesIO(file_content) as pdf_stream:
                pdf_reader = PyPDF2.PdfReader(pdf_stream)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF bytes: {str(e)}")
            return None
    
    def _extract_text_from_bytes_docx(self, file_content: bytes) -> Optional[str]:
        """Extract text from DOCX bytes"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX processing")
            return None
        
        try:
            import io
            text_content = []
            
            with io.BytesIO(file_content) as docx_stream:
                doc = docx.Document(docx_stream)
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX bytes: {str(e)}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove leading/trailing whitespace
            cleaned_line = line.strip()
            # Skip empty lines
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive consecutive newlines
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        
        return cleaned_text
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        path = Path(file_path)
        return path.suffix.lower() in self.supported_formats 