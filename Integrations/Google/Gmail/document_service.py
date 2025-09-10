"""
Document Service
Handles conversion of emails to Word documents and Google Drive storage
"""

import os
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import re

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try:
    from Integrations.Google.Gmail.gmail_service import EmailData
except ImportError:
    # Fallback imports with path adjustment
    import sys
    from pathlib import Path
    project_root = str(Path(__file__).parent.parent.parent.parent)
    if project_root not in sys.path:
        sys.path.append(project_root)
    
    from Integrations.Google.Gmail.gmail_service import EmailData

# Import Google Drive manager from Memory directory
from Integrations.Google.Drive.gdrive_manager import GoogleDriveManager

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for converting emails to Word documents and managing document storage"""
    
    def __init__(self, use_google_drive: bool = True, local_storage_path: str = None):
        """
        Initialize DocumentService
        
        Args:
            use_google_drive: Whether to use Google Drive for storage
            local_storage_path: Path for local storage (if not using Google Drive)
        """
        self.use_google_drive = use_google_drive
        self.local_storage_path = Path(local_storage_path) if local_storage_path else Path("AgentUI/Clients")
        
        # Ensure local storage directory exists
        self.local_storage_path.mkdir(parents=True, exist_ok=True)
        
        # Cache for Google Drive depth limit errors to avoid spam
        self._depth_limit_error_cache = set()
        
        logger.info(f"DocumentService initialized - Google Drive: {use_google_drive}, Local path: {self.local_storage_path}")
    
    def _get_google_drive_service(self, user_id: str):
        """Get Google Drive service using Gmail OAuth tokens"""
        try:
            from Integrations.Google.Gmail.gmail_database import gmail_database
            from Integrations.Google.Gmail.gmail_auth import GmailAuthService
            from googleapiclient.discovery import build
            
            # Get Gmail OAuth tokens
            token_data = gmail_database.get_gmail_tokens(user_id)
            if not token_data:
                logger.error("No Gmail tokens found for Google Drive access")
                return None
            
            # Use Gmail auth service to get credentials
            auth_service = GmailAuthService()
            credentials = auth_service.get_credentials(token_data, user_id)
            
            if not credentials:
                logger.error("Failed to get valid credentials for Google Drive")
                return None
            
            # Build Google Drive service
            drive_service = build('drive', 'v3', credentials=credentials)
            logger.info("Google Drive service initialized using Gmail OAuth")
            return drive_service
            
        except Exception as e:
            logger.error(f"Error initializing Google Drive service: {e}")
            return None
    
    def process_emails_to_documents(self, emails: List[EmailData], user_id: str = None) -> Dict[str, Any]:
        """
        Process a list of emails and convert them to Word documents
        
        Args:
            emails: List of EmailData objects to process
            
        Returns:
            Dictionary with processing results and statistics
        """
        try:
            import time
            start_time = time.time()
            total_emails = len(emails)
            
            logger.info(f"📄 Starting document processing for {total_emails} emails...")
            
            results = {
                'processed': 0,
                'failed': 0,
                'clients': {},
                'errors': []
            }
            
            for i, email in enumerate(emails, 1):
                try:
                    # Progress indicator
                    progress_msg = f"[{i}/{total_emails}]" if total_emails > 1 else ""
                    logger.info(f"📝 {progress_msg} Processing: {email.subject[:60]}...")
                    
                    document_start = time.time()
                    document_info = self.create_email_document(email, user_id)
                    document_time = time.time() - document_start
                    
                    if document_info:
                        client_name = email.client_info.get('name', 'Unknown')
                        
                        if client_name not in results['clients']:
                            results['clients'][client_name] = {
                                'count': 0,
                                'documents': []
                            }
                        
                        results['clients'][client_name]['count'] += 1
                        
                        # Create document entry with Google Drive or local information
                        doc_entry = {
                            'path': document_info['path'],
                            'subject': email.subject,
                            'timestamp': self._safe_timestamp_isoformat(email.timestamp),
                            'storage_type': document_info['storage_type']
                        }
                        
                        # Add Google Drive specific information if available
                        if document_info['storage_type'] == 'google_drive':
                            doc_entry['file_id'] = document_info.get('file_id')
                            doc_entry['web_link'] = document_info.get('web_link')
                        
                        results['clients'][client_name]['documents'].append(doc_entry)
                        
                        results['processed'] += 1
                        storage_type = document_info['storage_type']
                        if document_time > 3.0:  # Log slow operations
                            logger.info(f"✅ {progress_msg} Processed to {storage_type} ({document_time:.1f}s - slow): {email.subject[:50]}...")
                        else:
                            logger.info(f"✅ {progress_msg} Processed to {storage_type} ({document_time:.1f}s): {email.subject[:50]}...")
                    else:
                        results['failed'] += 1
                        logger.warning(f"⚠️ {progress_msg} Failed to create document: {email.subject[:50]}...")
                        
                except Exception as e:
                    error_msg = f"Error processing email {email.subject[:50]}: {str(e)}"
                    logger.error(f"❌ {progress_msg} {error_msg}")
                    # Check if it's a Google Drive depth limit error
                    if "myDriveHierarchyDepthLimitExceeded" in str(e) or "100 levels of folders" in str(e):
                        logger.warning(f"🚫 Google Drive folder depth limit reached - consider cleaning up your Drive")
                    results['errors'].append(error_msg)
                    results['failed'] += 1
            
            total_time = time.time() - start_time
            logger.info(f"🎉 Document processing complete ({total_time:.1f}s). Processed: {results['processed']}, Failed: {results['failed']}")
            
            # Summary by client
            if results['clients']:
                logger.info("📊 Processing summary by client:")
                for client_name, client_data in results['clients'].items():
                    logger.info(f"   • {client_name}: {client_data['count']} document(s)")
            
            return results
            
        except Exception as e:
            logger.error(f"Error in bulk email processing: {e}")
            return {
                'processed': 0,
                'failed': len(emails),
                'clients': {},
                'errors': [str(e)]
            }
    
    def create_email_document(self, email: EmailData, user_id: str = None) -> Optional[Dict[str, Any]]:
        """
        Create a Word document from an email and store it in Google Drive or locally
        
        Args:
            email: EmailData object to convert
            user_id: User ID for Google Drive access
            
        Returns:
            Dictionary with document information or None if failed
            Format: {
                'path': str,  # Google Drive path or local path
                'file_id': str,  # Google Drive file ID (if using Google Drive)
                'web_link': str,  # Google Drive web view link (if using Google Drive)
                'storage_type': str  # 'google_drive' or 'local'
                'is_duplicate': bool  # True if document already existed
            }
        """
        try:
            # Generate document filename
            document_name = self._generate_document_name(email)
            
            # Check for duplicates first
            if self.use_google_drive and user_id:
                existing_document = self._check_existing_document(email, document_name, user_id)
                if existing_document:
                    logger.info(f"📋 Document already exists: {document_name}")
                    existing_document['is_duplicate'] = True
                    return existing_document
            
            # Create structured Word document using new format
            try:
                from Integrations.Google.Gmail.structured_email_formatter import email_formatter
                
                # Prepare email data for structured format
                email_data = self._prepare_email_data_for_structured_format(email)
                
                # Create temporary file path for document creation
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_path = temp_file.name
                
                # Create structured document
                success = email_formatter.create_structured_email_document(email_data, temp_path)
                
                if success:
                    # Load the structured document for further processing
                    doc = Document(temp_path)
                    # Clean up temp file
                    os.unlink(temp_path)
                else:
                    # Fallback to old format
                    logger.warning("⚠️ Structured format creation failed, using legacy format")
                    doc = self._create_legacy_email_document(email)
                    
            except Exception as e:
                logger.error(f"❌ Error creating structured email document: {e}, falling back to legacy format")
                doc = self._create_legacy_email_document(email)
            
            if self.use_google_drive:
                # Save to Google Drive using Gmail OAuth
                result = self._save_to_google_drive(doc, email, document_name, user_id)
                if result:
                    result['is_duplicate'] = False
                return result
            else:
                # Save locally
                result = self._save_locally(doc, email, document_name)
                if result:
                    result['is_duplicate'] = False
                return result
            
        except Exception as e:
            logger.error(f"Error creating document for email {email.subject}: {e}")
            return None
    
    def _check_existing_document(self, email: EmailData, filename: str, user_id: str) -> Optional[Dict[str, Any]]:
        """
        Check if a document already exists in Google Drive
        
        Args:
            email: EmailData object
            filename: Document filename to check
            user_id: User ID for Google Drive access
            
        Returns:
            Existing document information if found, None otherwise
        """
        try:
            # Get Google Drive service
            drive_service = self._get_google_drive_service(user_id)
            if not drive_service:
                logger.warning("Google Drive not available for duplicate checking")
                return None
            
            # Get folder path from email categorization
            original_folder_path = email.client_info.get('folder_path', 'Uncategorized')
            
            # Build DigitalTwin_Brain folder structure
            digital_twin_folder_path = self._build_digital_twin_folder_path(original_folder_path, user_id)
            
            # Create folder structure (this will find existing folders)
            folder_id = self._create_google_drive_folder_path(drive_service, digital_twin_folder_path)
            if not folder_id:
                logger.warning(f"Failed to access Google Drive folder (likely due to depth limit): {digital_twin_folder_path}")
                logger.info("Skipping duplicate check due to Google Drive limitations")
                return None
            
            # Search for existing file with the same name in the folder
            query = f"name='{filename}' and '{folder_id}' in parents and trashed=false"
            
            results = drive_service.files().list(
                q=query,
                fields="files(id, name, webViewLink)",
                pageSize=10
            ).execute()
            
            files = results.get('files', [])
            
            if files:
                # Document already exists
                existing_file = files[0]  # Take the first match
                file_id = existing_file['id']
                web_link = existing_file.get('webViewLink', f"https://drive.google.com/file/d/{file_id}/view")
                gdrive_path = f"{digital_twin_folder_path}/{filename}"
                
                logger.info(f"📋 Found existing document: {gdrive_path}")
                return {
                    'path': gdrive_path,
                    'file_id': file_id,
                    'web_link': web_link,
                    'storage_type': 'google_drive'
                }
            
            # Document doesn't exist
            return None
            
        except Exception as e:
            logger.error(f"Error checking for existing document: {e}")
            return None
    
    def sync_database_with_google_drive(self, user_id: str, cleanup_orphaned: bool = True) -> Dict[str, Any]:
        """
        Synchronize database with Google Drive by removing entries for deleted documents
        
        Args:
            user_id: User ID to sync for
            cleanup_orphaned: Whether to remove database entries for missing Google Drive files
            
        Returns:
            Dictionary with sync results
        """
        sync_results = {
            'total_checked': 0,
            'existing_files': 0,
            'missing_files': 0,
            'removed_from_db': 0,
            'errors': []
        }
        
        try:
            # Get Google Drive service
            drive_service = self._get_google_drive_service(user_id)
            if not drive_service:
                sync_results['errors'].append("Could not access Google Drive")
                return sync_results
            
            # Get all processed emails for this user from database
            from .gmail_database import gmail_database
            
            if not gmail_database.supabase_client:
                sync_results['errors'].append("Could not connect to database")
                return sync_results
            
            # Query ALL processed emails (since document_path is not populated in current schema)
            result = gmail_database.supabase_client.table("processed_emails").select(
                "id, message_id, sender_email, subject, folder_path, client_name"
            ).eq("user_id", user_id).execute()
            
            if not result.data:
                logger.info(f"No processed emails found for user {user_id}")
                return sync_results
            
            sync_results['total_checked'] = len(result.data)
            
            # Check each document in Google Drive
            missing_entries = []
            
            for email_record in result.data:
                try:
                    original_folder_path = email_record.get('folder_path', 'Uncategorized')
                    message_id = email_record.get('message_id', '')
                    subject = email_record.get('subject', 'email')
                    email_id = email_record.get('id', 'unknown')
                    
                    # Build DigitalTwin_Brain folder structure for this user's emails
                    digital_twin_folder_path = self._build_digital_twin_folder_path(original_folder_path, user_id)
                    
                    # Construct expected filename based on current document creation pattern
                    # Since we don't store document_path, we'll check for likely filename patterns
                    safe_subject = self._sanitize_filename(subject)
                    
                    # Try multiple possible filename patterns that might have been used
                    possible_filenames = [
                        f"{safe_subject}.docx",
                        f"email-{email_id}-{safe_subject}.docx", 
                        f"{message_id}.docx",
                        f"{safe_subject[:30]}.docx"  # Truncated version
                    ]
                    
                    # Check if any of the possible files exist in Google Drive
                    file_exists = False
                    found_filename = None
                    
                    for filename in possible_filenames:
                        if self._check_google_drive_file_exists(drive_service, digital_twin_folder_path, filename):
                            file_exists = True
                            found_filename = filename
                            break
                    
                    if file_exists:
                        sync_results['existing_files'] += 1
                        logger.debug(f"✅ File exists: {digital_twin_folder_path}/{found_filename}")
                    else:
                        sync_results['missing_files'] += 1
                        missing_entries.append(email_record)
                        logger.info(f"❌ File missing: {digital_twin_folder_path}/<any-pattern>.docx (Message: {message_id})")
                
                except Exception as e:
                    logger.error(f"Error checking file for message {message_id}: {e}")
                    sync_results['errors'].append(f"Error checking {message_id}: {str(e)}")
            
            # Remove orphaned database entries
            if cleanup_orphaned and missing_entries:
                logger.info(f"Removing {len(missing_entries)} orphaned database entries...")
                
                for entry in missing_entries:
                    try:
                        # Remove the database entry
                        delete_result = gmail_database.supabase_client.table("processed_emails").delete().eq(
                            "id", entry['id']
                        ).execute()
                        
                        if delete_result.data:
                            sync_results['removed_from_db'] += 1
                            logger.info(f"🗑️ Removed orphaned entry: {entry.get('subject', 'Unknown')[:50]}")
                    
                    except Exception as e:
                        logger.error(f"Error removing orphaned entry {entry.get('id')}: {e}")
                        sync_results['errors'].append(f"Error removing entry {entry.get('id')}: {str(e)}")
            
            # Log summary
            logger.info(f"📊 Google Drive-Database Sync Complete:")
            logger.info(f"   Total checked: {sync_results['total_checked']}")
            logger.info(f"   Files exist: {sync_results['existing_files']}")
            logger.info(f"   Files missing: {sync_results['missing_files']}")
            logger.info(f"   Removed from DB: {sync_results['removed_from_db']}")
            
            return sync_results
            
        except Exception as e:
            logger.error(f"Error in Google Drive-Database sync: {e}")
            sync_results['errors'].append(f"Sync error: {str(e)}")
            return sync_results
    
    def _check_google_drive_file_exists(self, drive_service, folder_path: str, filename: str) -> bool:
        """
        Check if a specific file exists in Google Drive
        
        Args:
            drive_service: Google Drive service instance
            folder_path: Folder path (e.g., 'Internal/Repsmate Team')
            filename: File name to check
            
        Returns:
            True if file exists, False otherwise
        """
        try:
            # Get or create the folder structure
            folder_id = self._create_google_drive_folder_path(drive_service, folder_path)
            if not folder_id:
                return False
            
            # Escape special characters in filename for Google Drive query
            escaped_filename = self._escape_google_drive_filename(filename)
            
            # Search for the file in the specific folder
            query = f"name='{escaped_filename}' and '{folder_id}' in parents and trashed=false"
            
            results = drive_service.files().list(
                q=query,
                fields="files(id, name)",
                pageSize=1
            ).execute()
            
            files = results.get('files', [])
            return len(files) > 0
            
        except Exception as e:
            logger.error(f"Error checking Google Drive file existence for '{filename}': {e}")
            return False
    
    def _escape_google_drive_filename(self, filename: str) -> str:
        """
        Escape special characters in filename for Google Drive API query
        
        Google Drive search requires escaping of: \\ ' " 
        """
        # Escape backslashes first (must be first to avoid double-escaping)
        escaped = filename.replace('\\', '\\\\')
        # Escape single quotes
        escaped = escaped.replace("'", "\\'")
        # Escape double quotes  
        escaped = escaped.replace('"', '\\"')
        
        return escaped
    
    def _sanitize_filename(self, text: str) -> str:
        """Sanitize text for use as filename and Google Drive queries"""
        import re
        # Remove or replace special characters that cause Google Drive API issues
        # Replace problematic characters with safe alternatives
        sanitized = text.replace("'", "").replace('"', '').replace('[', '').replace(']', '')
        sanitized = sanitized.replace('(', '').replace(')', '').replace('&', 'and')
        sanitized = sanitized.replace('#', '').replace('%', '').replace('@', 'at')
        sanitized = sanitized.replace('+', 'plus').replace('=', '').replace('!', '')
        sanitized = sanitized.replace('~', '').replace('`', '').replace('{', '').replace('}', '')
        
        # Remove remaining problematic characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '', sanitized)
        
        # Replace multiple spaces with single hyphens
        sanitized = re.sub(r'\s+', '-', sanitized)
        
        # Remove multiple consecutive hyphens
        sanitized = re.sub(r'-+', '-', sanitized)
        
        # Trim hyphens from start and end
        sanitized = sanitized.strip('-')
        
        # Ensure filename is not empty and limit length
        if not sanitized:
            sanitized = 'email'
        
        return sanitized[:50] if len(sanitized) > 50 else sanitized
    
    def _safe_timestamp_isoformat(self, timestamp) -> Optional[str]:
        """
        Safely convert timestamp to ISO format string
        
        Args:
            timestamp: Either a datetime object or ISO string
            
        Returns:
            ISO format string or None
        """
        if not timestamp:
            return None
            
        try:
            if isinstance(timestamp, str):
                # Try to parse and re-format to ensure consistency
                try:
                    # Handle ISO strings with timezone
                    timestamp_obj = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                    return timestamp_obj.isoformat()
                except (ValueError, AttributeError):
                    # Try parsing with email.utils for other formats
                    import email.utils
                    try:
                        timestamp_obj = email.utils.parsedate_to_datetime(timestamp)
                        return timestamp_obj.isoformat()
                    except:
                        # Return the string as-is if we can't parse it
                        return timestamp
            elif hasattr(timestamp, 'isoformat'):
                # Already a datetime object
                return timestamp.isoformat()
            else:
                # Unknown type, try to convert to string
                return str(timestamp)
                
        except Exception as e:
            logger.warning(f"Error converting timestamp to ISO format: {e}")
            return str(timestamp) if timestamp else None
    
    def _save_to_google_drive(self, doc: Document, email: EmailData, filename: str, user_id: str = None) -> Optional[Dict[str, Any]]:
        """Save document to Google Drive using Gmail OAuth"""
        try:
            # Get user_id from email processing context if not provided
            if not user_id:
                # Try to get user_id from email metadata or use a default approach
                # For now, we'll need to pass user_id from the calling method
                logger.error("User ID required for Google Drive access")
                return self._save_locally(doc, email, filename)
            
            # Get Google Drive service
            drive_service = self._get_google_drive_service(user_id)
            if not drive_service:
                logger.warning("Google Drive not available, falling back to local storage")
                return self._save_locally(doc, email, filename)
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = temp_file.name
                doc.save(temp_path)
            
            # Get folder path from email categorization
            original_folder_path = email.client_info.get('folder_path', 'Uncategorized')
            
            # Build DigitalTwin_Brain folder structure
            digital_twin_folder_path = self._build_digital_twin_folder_path(original_folder_path, user_id)
            
            # Create folder structure in Google Drive
            folder_id = self._create_google_drive_folder_path(drive_service, digital_twin_folder_path)
            
            if not folder_id:
                logger.warning(f"Failed to create Google Drive folder (likely due to depth limit): {digital_twin_folder_path}")
                logger.info("Skipping document creation due to Google Drive limitations")
                # Clean up temp file
                os.unlink(temp_path)
                # Don't fall back to local storage to avoid cluttering - just skip this email
                return None
            
            # Upload file to Google Drive
            file_id = self._upload_to_google_drive(drive_service, temp_path, folder_id, filename)
            
            # Clean up temporary file
            os.unlink(temp_path)
            
            if file_id:
                # Generate web link
                web_link = f"https://drive.google.com/file/d/{file_id}/view"
                gdrive_path = f"{digital_twin_folder_path}/{filename}"
                
                logger.info(f"Created document in Google Drive: {gdrive_path}")
                return {
                    'path': gdrive_path,
                    'file_id': file_id,
                    'web_link': web_link,
                    'storage_type': 'google_drive'
                }
            else:
                logger.error(f"Failed to upload document to Google Drive: {filename}")
                return self._save_locally(doc, email, filename)
                
        except Exception as e:
            logger.error(f"Error saving document to Google Drive: {e}")
            # Clean up temp file if it exists
            try:
                if 'temp_path' in locals():
                    os.unlink(temp_path)
            except:
                pass
            # Fall back to local storage
            return self._save_locally(doc, email, filename)
    
    def _create_google_drive_folder_path(self, drive_service, folder_path: str, root_folder_id: str = None) -> Optional[str]:
        """Create nested folder structure in Google Drive"""
        try:
            if not folder_path:
                return root_folder_id
            
            # Check if we've already failed on this path due to depth limits
            if folder_path in self._depth_limit_error_cache:
                return None  # Silently fail for already-known problematic paths
            
            folder_parts = [part.strip() for part in folder_path.split('/') if part.strip()]
            current_parent_id = root_folder_id
            
            for folder_name in folder_parts:
                folder_id = self._find_or_create_google_drive_folder(drive_service, folder_name, current_parent_id)
                if not folder_id:
                    # If any folder creation fails, cache the full path to avoid retrying
                    if "depth limit" in str(folder_name).lower():  # This is a rough check
                        self._depth_limit_error_cache.add(folder_path)
                    logger.error(f"Failed to create/find folder: {folder_name}")
                    return None
                current_parent_id = folder_id
            
            return current_parent_id
            
        except Exception as e:
            # If there's a depth limit error, cache this path to avoid retrying
            if "myDriveHierarchyDepthLimitExceeded" in str(e) or "100 levels of folders" in str(e):
                self._depth_limit_error_cache.add(folder_path)
            logger.error(f"Error creating folder path: {e}")
            return None
    
    def _find_or_create_google_drive_folder(self, drive_service, name: str, parent_id: str = None) -> Optional[str]:
        """Find existing folder or create new one in Google Drive"""
        try:
            # Search for existing folder
            query = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
            if parent_id:
                query += f" and '{parent_id}' in parents"
            
            results = drive_service.files().list(q=query, fields="files(id, name)").execute()
            folders = results.get('files', [])
            
            if folders:
                folder_id = folders[0]['id']
                logger.info(f"Found existing Google Drive folder '{name}' with ID: {folder_id}")
                return folder_id
            else:
                # Create new folder
                return self._create_google_drive_folder(drive_service, name, parent_id)
                
        except Exception as e:
            logger.error(f"Error finding/creating Google Drive folder '{name}': {e}")
            return None
    
    def _create_google_drive_folder(self, drive_service, name: str, parent_id: str = None) -> Optional[str]:
        """Create a folder in Google Drive"""
        try:
            folder_metadata = {
                'name': name,
                'mimeType': 'application/vnd.google-apps.folder'
            }
            
            if parent_id:
                folder_metadata['parents'] = [parent_id]
            
            folder = drive_service.files().create(body=folder_metadata, fields='id').execute()
            folder_id = folder.get('id')
            
            logger.info(f"Created Google Drive folder '{name}' with ID: {folder_id}")
            return folder_id
            
        except Exception as e:
            # Check if it's a depth limit error
            if "myDriveHierarchyDepthLimitExceeded" in str(e) or "100 levels of folders" in str(e):
                # Only log the first time we encounter this error
                error_key = f"depth_limit_{name}_{parent_id}"
                if error_key not in self._depth_limit_error_cache:
                    logger.warning(f"🚫 Google Drive depth limit reached for folder '{name}'")
                    logger.warning("💡 Consider cleaning up your Google Drive folder structure to reduce nesting")
                    self._depth_limit_error_cache.add(error_key)
                return None  # Return None to trigger graceful fallback
            else:
                logger.error(f"Error creating Google Drive folder '{name}': {e}")
                return None
    
    def _upload_to_google_drive(self, drive_service, file_path: str, folder_id: str, filename: str) -> Optional[str]:
        """Upload a file to Google Drive"""
        try:
            from googleapiclient.http import MediaFileUpload
            
            file_metadata = {'name': filename}
            if folder_id:
                file_metadata['parents'] = [folder_id]
            
            media = MediaFileUpload(
                file_path, 
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            file = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id,name,webViewLink'
            ).execute()
            
            file_id = file.get('id')
            logger.info(f"Uploaded file '{filename}' to Google Drive with ID: {file_id}")
            
            return file_id
            
        except Exception as e:
            logger.error(f"Error uploading file to Google Drive: {e}")
            return None
    
    def _save_locally(self, doc: Document, email: EmailData, filename: str) -> Optional[Dict[str, Any]]:
        """Save document locally (fallback)"""
        try:
            # Get folder using categorization system
            folder_path = email.client_info.get('folder_path', 'Uncategorized')
            client_folder = self._get_folder_by_path(folder_path)
            
            # Save document
            document_path = client_folder / filename
            doc.save(str(document_path))
            
            logger.info(f"Created local document: {document_path}")
            return {
                'path': str(document_path),
                'file_id': None,
                'web_link': None,
                'storage_type': 'local'
            }
            
        except Exception as e:
            logger.error(f"Error saving document locally: {e}")
            return None
    
    def _get_client_folder(self, client_name: str) -> Path:
        """
        Get or create client-specific folder (legacy method)
        
        Args:
            client_name: Name of the client
            
        Returns:
            Path to client folder
        """
        # Sanitize client name for filesystem
        safe_client_name = self._sanitize_filename(client_name)
        client_folder = self.local_storage_path / safe_client_name
        client_folder.mkdir(exist_ok=True)
        return client_folder
    
    def _get_folder_by_path(self, folder_path: str) -> Path:
        """
        Get or create folder using hierarchical path (e.g., 'Internal/Repsamte Team', 'Clients/Amazon')
        
        Args:
            folder_path: Hierarchical folder path (e.g., 'Internal/Repsamte Team')
            
        Returns:
            Path to the folder
        """
        # Split path and sanitize each component
        path_parts = [self._sanitize_filename(part.strip()) for part in folder_path.split('/') if part.strip()]
        
        # Build the full path
        full_path = self.local_storage_path
        for part in path_parts:
            full_path = full_path / part
        
        # Create the directory structure
        full_path.mkdir(parents=True, exist_ok=True)
        return full_path
    
    def _generate_document_name(self, email: EmailData) -> str:
        """
        Generate a safe filename for the email document using format: {email-subject}-{sender_name}
        
        Args:
            email: EmailData object
            
        Returns:
            Safe filename string
        """
        try:
            # Sanitize subject for filename (limit to 50 chars for readability)
            safe_subject = self._sanitize_filename(email.subject[:50])
            if not safe_subject:
                safe_subject = "no-subject"
            
            # Extract sender name from sender_email (the part before @)
            sender_name = ""
            if email.sender_email:
                # Get the part before @ and sanitize it
                sender_prefix = email.sender_email.split('@')[0]
                sender_name = self._sanitize_filename(sender_prefix)
            
            # If no sender name available, use sender field or fallback
            if not sender_name:
                if email.sender:
                    # Extract name from sender field (might be in format "Name <email@domain.com>")
                    sender_clean = email.sender.split('<')[0].strip()
                    sender_name = self._sanitize_filename(sender_clean)
                else:
                    sender_name = "unknown-sender"
            
            # Ensure sender name is not empty and limit length
            if not sender_name:
                sender_name = "unknown-sender"
            else:
                sender_name = sender_name[:30]  # Limit sender name length
            
            # Create filename in format: {email-subject}-{sender_name}
            filename = f"{safe_subject}-{sender_name}.docx"
            
            # Ensure total filename length is reasonable (max 100 chars before extension)
            base_name = filename.replace('.docx', '')
            if len(base_name) > 100:
                # Trim subject if needed, keeping sender name
                max_subject_len = 100 - len(sender_name) - 1  # -1 for the hyphen
                if max_subject_len > 10:  # Ensure we have at least some subject
                    safe_subject = safe_subject[:max_subject_len]
                    filename = f"{safe_subject}-{sender_name}.docx"
                else:
                    # If sender name is too long, trim both
                    safe_subject = safe_subject[:50]
                    sender_name = sender_name[:30]
                    filename = f"{safe_subject}-{sender_name}.docx"
            
            return filename
            
        except Exception as e:
            logger.error(f"Error generating document name: {e}")
            # Fallback filename
            safe_subject = self._sanitize_filename(email.subject[:30]) if email.subject else "email"
            return f"{safe_subject}-fallback-{email.message_id[:8]}.docx"
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize string for use as filename and Google Drive queries
        
        Args:
            filename: Raw filename string
            
        Returns:
            Sanitized filename string
        """
        # Remove or replace special characters that cause Google Drive API issues
        # Replace problematic characters with safe alternatives
        sanitized = filename.replace("'", "").replace('"', '').replace('[', '').replace(']', '')
        sanitized = sanitized.replace('(', '').replace(')', '').replace('&', 'and')
        sanitized = sanitized.replace('#', '').replace('%', '').replace('@', 'at')
        sanitized = sanitized.replace('+', 'plus').replace('=', '').replace('!', '')
        sanitized = sanitized.replace('~', '').replace('`', '').replace('{', '').replace('}', '')
        
        # Remove remaining problematic characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '', sanitized)
        
        # Replace multiple spaces with single hyphens
        sanitized = re.sub(r'\s+', '-', sanitized)
        
        # Remove multiple consecutive hyphens
        sanitized = re.sub(r'-+', '-', sanitized)
        
        # Trim hyphens from start and end
        sanitized = sanitized.strip('-')
        
        # Ensure filename is not empty and limit length
        if not sanitized:
            sanitized = 'email'
        
        return sanitized[:50] if len(sanitized) > 50 else sanitized
    
    def _get_user_info_from_tokens(self, user_id: str) -> Dict[str, str]:
        """
        Extract user information from Gmail tokens to get username and email
        
        Args:
            user_id: User identifier
            
        Returns:
            Dictionary with user info (username, email, name)
        """
        try:
            from .gmail_database import gmail_database
            
            # Get user's Gmail tokens which contain user info
            token_data = gmail_database.get_gmail_tokens(user_id)
            if not token_data or 'user_info' not in token_data:
                logger.warning(f"No user info found for user {user_id}")
                return {
                    'username': user_id,  # Fallback to user_id
                    'email': 'unknown@example.com',
                    'name': 'Unknown User'
                }
            
            user_info = token_data['user_info']
            email = user_info.get('email', 'unknown@example.com')
            name = user_info.get('name', 'Unknown User')
            
            # Extract username from email (part before @)
            username = email.split('@')[0] if '@' in email else user_id
            
            # Sanitize username for folder name
            safe_username = self._sanitize_filename(username)
            
            logger.info(f"Extracted user info for {user_id}: username={safe_username}, email={email}")
            
            return {
                'username': safe_username,
                'email': email,
                'name': name
            }
            
        except Exception as e:
            logger.error(f"Error extracting user info for {user_id}: {e}")
            return {
                'username': user_id,
                'email': 'unknown@example.com',
                'name': 'Unknown User'
            }
    
    def _build_digital_twin_folder_path(self, original_folder_path: str, user_id: str) -> str:
        """
        Build the complete DigitalTwin_Brain folder path for a user's emails
        
        Args:
            original_folder_path: Original folder path (e.g., 'Internal/Repsmate Team')
            user_id: User identifier
            
        Returns:
            Complete folder path with DigitalTwin_Brain structure
            Format: DigitalTwin_Brain/Users/{username}/Emails/{original_path}
        """
        try:
            # Check if path already has DigitalTwin_Brain structure (from semantic pipeline)
            if original_folder_path.startswith('DigitalTwin_Brain/'):
                logger.debug(f"Path already has DigitalTwin_Brain structure: {original_folder_path}")
                return original_folder_path
            
            user_info = self._get_user_info_from_tokens(user_id)
            username = user_info['username']
            
            # Build the complete path: DigitalTwin_Brain/Users/{username}/Emails/{original_path}
            complete_path = f"DigitalTwin_Brain/Users/{username}/Emails/{original_folder_path}"
            
            logger.debug(f"Built DigitalTwin folder path: {complete_path}")
            return complete_path
            
        except Exception as e:
            logger.error(f"Error building DigitalTwin folder path: {e}")
            # Fallback to original path with generic structure
            return f"DigitalTwin_Brain/Users/{user_id}/Emails/{original_folder_path}"
    
    def _add_metadata_section(self, doc: Document, email: EmailData) -> None:
        """
        Add metadata section to document
        
        Args:
            doc: Document object
            email: EmailData object
        """
        try:
            # Add metadata heading
            doc.add_heading('Email Details', level=1)
            
            # Create metadata table
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            # Safely get email attributes
            sender = getattr(email, 'sender', '') or getattr(email, 'sender_email', '') or 'Unknown'
            sender_email = getattr(email, 'sender_email', '') or 'Unknown'
            subject = getattr(email, 'subject', '') or 'No Subject'
            message_id = getattr(email, 'message_id', '') or 'Unknown'
            thread_id = getattr(email, 'thread_id', '') or 'Unknown'
            
            # Handle timestamp safely
            timestamp_str = 'Unknown'
            if hasattr(email, 'timestamp') and email.timestamp:
                try:
                    if isinstance(email.timestamp, str):
                        timestamp_str = email.timestamp
                    else:
                        timestamp_str = email.timestamp.strftime('%Y-%m-%d %H:%M:%S')
                except:
                    timestamp_str = str(email.timestamp)
            
            # Add metadata rows
            metadata_items = [
                ('Subject', subject),
                ('From', sender),
                ('Sender Email', sender_email),
                ('Date', timestamp_str),
                ('Message ID', message_id),
                ('Thread ID', thread_id)
            ]
            
            for label, value in metadata_items:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value) if value else 'N/A'
                
                # Make label cell bold
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add spacing
            doc.add_paragraph()
            
        except Exception as e:
            logger.error(f"Error adding metadata section: {e}")
    
    def _add_content_section(self, doc: Document, email: EmailData) -> None:
        """
        Add email content section to document
        
        Args:
            doc: Document object
            email: EmailData object
        """
        try:
            # Add content heading
            doc.add_heading('Email Content', level=1)
            
            # Add email body
            if email.body:
                # Split body into paragraphs
                paragraphs = email.body.split('\n\n')
                
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        paragraph = doc.add_paragraph(para_text)
                        paragraph.space_after = Inches(0.1)
            else:
                doc.add_paragraph('No content available.')
            
            # Add spacing
            doc.add_paragraph()
            
        except Exception as e:
            logger.error(f"Error adding content section: {e}")
    
    def _prepare_email_data_for_structured_format(self, email: EmailData) -> Dict[str, Any]:
        """
        Prepare email data for the new structured document format
        
        Args:
            email: EmailData object
            
        Returns:
            Dictionary formatted for structured email formatter
        """
        try:
            # Extract sender information
            sender_info = {
                'name': getattr(email, 'sender', '') or 'Unknown Sender',
                'email': getattr(email, 'sender_email', '') or ''
            }
            
            # Format timestamp
            timestamp_str = ''
            if hasattr(email, 'timestamp') and email.timestamp:
                try:
                    if isinstance(email.timestamp, str):
                        timestamp_str = email.timestamp
                    else:
                        timestamp_str = email.timestamp.strftime('%Y-%m-%d %H:%M:%S')
                except:
                    timestamp_str = str(email.timestamp)
            
            # Prepare recipients
            recipients = []
            if hasattr(email, 'recipients') and email.recipients:
                recipients = email.recipients if isinstance(email.recipients, list) else [email.recipients]
            
            # Determine category
            category = self._determine_email_category(email)
            
            # Prepare structured data
            email_data = {
                'subject': getattr(email, 'subject', '') or 'No Subject',
                'sender': sender_info,
                'date': timestamp_str,
                'recipients': recipients,
                'category': category,
                'body': getattr(email, 'body', '') or '',
                'thread_id': getattr(email, 'thread_id', ''),
                'message_id': getattr(email, 'message_id', ''),
                'labels': getattr(email, 'labels', []) or []
            }
            
            # Add client information if available
            if hasattr(email, 'client_info') and email.client_info.get('name') != 'Unknown':
                email_data['client_info'] = email.client_info
            
            return email_data
            
        except Exception as e:
            logger.error(f"Error preparing email data for structured format: {e}")
            return {
                'subject': 'Error Processing Email',
                'sender': {'name': 'Unknown', 'email': ''},
                'body': 'Error occurred while processing email data'
            }
    
    def _determine_email_category(self, email: EmailData) -> str:
        """
        Determine email category based on content and metadata
        
        Args:
            email: EmailData object
            
        Returns:
            Category string
        """
        try:
            subject = getattr(email, 'subject', '').lower()
            sender_email = getattr(email, 'sender_email', '').lower()
            body = getattr(email, 'body', '').lower()
            
            # Internal email detection
            internal_domains = ['repsmate.com', 'company.com']  # Add your domains
            if any(domain in sender_email for domain in internal_domains):
                return 'Internal'
            
            # Business category keywords
            business_keywords = ['contract', 'proposal', 'invoice', 'payment', 'business', 'deal', 'partnership']
            if any(keyword in subject or keyword in body for keyword in business_keywords):
                return 'Business'
            
            # Meeting category
            meeting_keywords = ['meeting', 'call', 'schedule', 'appointment', 'calendar']
            if any(keyword in subject or keyword in body for keyword in meeting_keywords):
                return 'Meeting'
            
            # Client category
            client_keywords = ['client', 'customer', 'support', 'issue', 'help']
            if any(keyword in subject or keyword in body for keyword in client_keywords):
                return 'Client'
            
            return 'General'
            
        except Exception as e:
            logger.error(f"Error determining email category: {e}")
            return 'General'
    
    def _create_legacy_email_document(self, email: EmailData) -> Document:
        """
        Create email document using the legacy format (fallback)
        
        Args:
            email: EmailData object
            
        Returns:
            Document object
        """
        try:
            doc = Document()
            
            # Add document title
            title = doc.add_heading(f'Email: {email.subject}', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata section
            self._add_metadata_section(doc, email)
            
            # Add email content
            self._add_content_section(doc, email)
            
            # Add client information
            if hasattr(email, 'client_info') and email.client_info.get('name') != 'Unknown':
                self._add_client_section(doc, email.client_info)
            
            return doc
            
        except Exception as e:
            logger.error(f"Error creating legacy email document: {e}")
            # Return minimal document
            doc = Document()
            doc.add_heading('Email Processing Error', 0)
            doc.add_paragraph(f"Error: {str(e)}")
            return doc
    
    def _add_client_section(self, doc: Document, client_info: Dict[str, Any]) -> None:
        """
        Add client information section to document
        
        Args:
            doc: Document object
            client_info: Client information dictionary
        """
        try:
            # Add client heading
            doc.add_heading('Client Information', level=1)
            
            # Create client info table
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            # Add client info rows
            client_items = [
                ('Client Name', client_info.get('name', 'Unknown')),
                ('Domain', client_info.get('domain', 'Unknown')),
                ('Detection Confidence', client_info.get('confidence', 'Unknown')),
                ('Detection Method', client_info.get('detection_method', 'Unknown'))
            ]
            
            for label, value in client_items:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value)
                
                # Make label cell bold
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
        except Exception as e:
            logger.error(f"Error adding client section: {e}")
    
    def get_client_summary(self) -> Dict[str, Any]:
        """
        Get summary of stored client documents
        
        Returns:
            Dictionary with client document statistics
        """
        try:
            summary = {
                'total_clients': 0,
                'total_documents': 0,
                'clients': {}
            }
            
            if not self.local_storage_path.exists():
                return summary
            
            for client_folder in self.local_storage_path.iterdir():
                if client_folder.is_dir():
                    client_name = client_folder.name
                    documents = list(client_folder.glob('*.docx'))
                    document_count = len(documents)
                    
                    if document_count > 0:
                        summary['clients'][client_name] = {
                            'document_count': document_count,
                            'folder_path': str(client_folder),
                            'latest_document': self._get_latest_document_info(documents)
                        }
                        
                        summary['total_clients'] += 1
                        summary['total_documents'] += document_count
            
            return summary
            
        except Exception as e:
            logger.error(f"Error getting client summary: {e}")
            return {'total_clients': 0, 'total_documents': 0, 'clients': {}}
    
    def _get_latest_document_info(self, documents: List[Path]) -> Dict[str, Any]:
        """
        Get information about the latest document in a list
        
        Args:
            documents: List of document paths
            
        Returns:
            Dictionary with latest document info
        """
        try:
            if not documents:
                return {}
            
            # Sort by modification time
            latest_doc = max(documents, key=lambda x: x.stat().st_mtime)
            mod_time = datetime.fromtimestamp(latest_doc.stat().st_mtime)
            
            return {
                'name': latest_doc.name,
                'path': str(latest_doc),
                'modified': mod_time.isoformat(),
                'size_bytes': latest_doc.stat().st_size
            }
            
        except Exception as e:
            logger.error(f"Error getting latest document info: {e}")
            return {}
    
    def cleanup_old_documents(self, days_old: int = 90) -> Dict[str, Any]:
        """
        Clean up documents older than specified days
        
        Args:
            days_old: Number of days after which to consider documents old
            
        Returns:
            Dictionary with cleanup results
        """
        try:
            cutoff_time = datetime.now().timestamp() - (days_old * 24 * 60 * 60)
            results = {
                'deleted_count': 0,
                'deleted_size_bytes': 0,
                'errors': []
            }
            
            for client_folder in self.local_storage_path.iterdir():
                if client_folder.is_dir():
                    for document in client_folder.glob('*.docx'):
                        try:
                            if document.stat().st_mtime < cutoff_time:
                                size_bytes = document.stat().st_size
                                document.unlink()
                                results['deleted_count'] += 1
                                results['deleted_size_bytes'] += size_bytes
                                logger.info(f"Deleted old document: {document}")
                        except Exception as e:
                            error_msg = f"Error deleting {document}: {str(e)}"
                            logger.error(error_msg)
                            results['errors'].append(error_msg)
            
            logger.info(f"Cleanup complete. Deleted {results['deleted_count']} documents")
            return results
            
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")
            return {'deleted_count': 0, 'deleted_size_bytes': 0, 'errors': [str(e)]}


# Create global instance
document_service = DocumentService() 