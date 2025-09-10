#!/usr/bin/env python3
"""
Structured Email Document Formatter
===================================

Creates structured, semantic-search optimized .docx files for emails
with standardized metadata and content layout.
"""

import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class EmailDocumentFormatter:
    """Creates structured email documents optimized for semantic search"""
    
    def __init__(self):
        self.style_applied = False
    
    def create_structured_email_document(self, email_data: Dict[str, Any], output_path: str) -> bool:
        """
        Create a structured .docx email document
        
        Args:
            email_data: Dictionary containing email information
            output_path: Path where to save the .docx file
            
        Returns:
            True if successful, False otherwise
        """
        if not DOCX_AVAILABLE:
            logger.error("❌ python-docx not available. Install with: pip install python-docx")
            return False
        
        try:
            doc = Document()
            
            # Apply document styling
            self._apply_document_styles(doc)
            
            # Add structured metadata header
            self._add_metadata_header(doc, email_data)
            
            # Add separator
            self._add_separator(doc)
            
            # Add conversation content
            self._add_conversation_content(doc, email_data)
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Save document
            doc.save(output_path)
            logger.info(f"✅ Structured email document created: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error creating structured email document: {e}")
            return False
    
    def _apply_document_styles(self, doc: Document):
        """Apply consistent styling to the document"""
        try:
            # Create or modify styles for better readability
            styles = doc.styles
            
            # Header style for metadata
            if 'EmailHeader' not in [s.name for s in styles]:
                header_style = styles.add_style('EmailHeader', WD_STYLE_TYPE.PARAGRAPH)
                header_style.font.name = 'Calibri'
                header_style.font.size = Pt(11)
                header_style.font.bold = True
            
            # Body style for content
            if 'EmailBody' not in [s.name for s in styles]:
                body_style = styles.add_style('EmailBody', WD_STYLE_TYPE.PARAGRAPH)
                body_style.font.name = 'Calibri'
                body_style.font.size = Pt(10)
            
            # Timestamp style
            if 'EmailTimestamp' not in [s.name for s in styles]:
                timestamp_style = styles.add_style('EmailTimestamp', WD_STYLE_TYPE.PARAGRAPH)
                timestamp_style.font.name = 'Calibri'
                timestamp_style.font.size = Pt(9)
                timestamp_style.font.italic = True
            
            self.style_applied = True
            
        except Exception as e:
            logger.warning(f"⚠️ Could not apply custom styles: {e}")
            self.style_applied = False
    
    def _add_metadata_header(self, doc: Document, email_data: Dict[str, Any]):
        """Add structured metadata header"""
        # Subject
        subject = email_data.get('subject', 'No Subject')
        p = doc.add_paragraph()
        p.add_run('Subject: ').bold = True
        p.add_run(subject)
        if self.style_applied:
            p.style = 'EmailHeader'
        
        # Sender
        sender = email_data.get('sender', {})
        sender_name = sender.get('name', 'Unknown Sender')
        sender_email = sender.get('email', '')
        
        p = doc.add_paragraph()
        p.add_run('Sender: ').bold = True
        if sender_email:
            p.add_run(f"{sender_name} <{sender_email}>")
        else:
            p.add_run(sender_name)
        if self.style_applied:
            p.style = 'EmailHeader'
        
        # Date
        date = email_data.get('date', email_data.get('timestamp', ''))
        if date:
            # Ensure date is in readable format
            formatted_date = self._format_date(date)
            p = doc.add_paragraph()
            p.add_run('Date: ').bold = True
            p.add_run(formatted_date)
            if self.style_applied:
                p.style = 'EmailHeader'
        
        # Recipients
        recipients = email_data.get('recipients', [])
        if recipients:
            p = doc.add_paragraph()
            p.add_run('To: ').bold = True
            if isinstance(recipients, list):
                p.add_run(', '.join(recipients))
            else:
                p.add_run(str(recipients))
            if self.style_applied:
                p.style = 'EmailHeader'
        
        # Category (if available)
        category = email_data.get('category', self._infer_category(email_data))
        if category:
            p = doc.add_paragraph()
            p.add_run('Category: ').bold = True
            p.add_run(category)
            if self.style_applied:
                p.style = 'EmailHeader'
        
        # Thread ID (if available)
        thread_id = email_data.get('thread_id', '')
        if thread_id:
            p = doc.add_paragraph()
            p.add_run('Thread ID: ').bold = True
            p.add_run(thread_id)
            if self.style_applied:
                p.style = 'EmailHeader'
        
        # Labels/Tags (if available)
        labels = email_data.get('labels', [])
        if labels:
            p = doc.add_paragraph()
            p.add_run('Labels: ').bold = True
            p.add_run(', '.join(labels))
            if self.style_applied:
                p.style = 'EmailHeader'
    
    def _add_separator(self, doc: Document):
        """Add a visual separator between metadata and content"""
        doc.add_paragraph()
        p = doc.add_paragraph('─' * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Add conversation header
        p = doc.add_paragraph()
        p.add_run('Conversation:').bold = True
        if self.style_applied:
            p.style = 'EmailHeader'
        doc.add_paragraph()
    
    def _add_conversation_content(self, doc: Document, email_data: Dict[str, Any]):
        """Add the conversation content in a structured format"""
        # Get conversation messages
        messages = email_data.get('messages', [])
        body = email_data.get('body', '')
        
        if messages:
            # Multi-message conversation
            for i, message in enumerate(messages):
                self._add_message(doc, message, i)
        elif body:
            # Single message
            self._add_single_message(doc, email_data)
        else:
            # Fallback content
            p = doc.add_paragraph("No content available")
            if self.style_applied:
                p.style = 'EmailBody'
    
    def _add_message(self, doc: Document, message: Dict[str, Any], index: int):
        """Add a single message to the conversation"""
        # Message timestamp and sender
        timestamp = message.get('timestamp', message.get('date', ''))
        sender = message.get('sender', message.get('from', 'Unknown'))
        
        if timestamp and sender:
            formatted_time = self._format_timestamp(timestamp)
            p = doc.add_paragraph(f"[{formatted_time}] {sender}:")
            p.style = 'EmailTimestamp' if self.style_applied else None
        elif sender:
            p = doc.add_paragraph(f"{sender}:")
            p.style = 'EmailTimestamp' if self.style_applied else None
        
        # Message content
        content = message.get('content', message.get('body', ''))
        if content:
            # Clean and format content
            cleaned_content = self._clean_content(content)
            p = doc.add_paragraph(cleaned_content)
            if self.style_applied:
                p.style = 'EmailBody'
        
        # Add spacing between messages
        if index < len(message) - 1:  # Not the last message
            doc.add_paragraph()
    
    def _add_single_message(self, doc: Document, email_data: Dict[str, Any]):
        """Add content for a single message email"""
        body = email_data.get('body', '')
        
        if body:
            # Clean and format the body
            cleaned_body = self._clean_content(body)
            
            # Check if the body already contains timestamp/sender info
            if self._has_conversation_format(cleaned_body):
                # Body already has conversation format, use as-is
                lines = cleaned_body.split('\n')
                for line in lines:
                    if line.strip():
                        p = doc.add_paragraph(line.strip())
                        if self.style_applied:
                            # Try to detect if it's a timestamp line
                            if re.match(r'\[.*\].*:', line.strip()):
                                p.style = 'EmailTimestamp'
                            else:
                                p.style = 'EmailBody'
            else:
                # Simple body content
                p = doc.add_paragraph(cleaned_body)
                if self.style_applied:
                    p.style = 'EmailBody'
    
    def _format_date(self, date_str: str) -> str:
        """Format date string to a consistent format"""
        try:
            # Try to parse various date formats
            if isinstance(date_str, datetime):
                return date_str.strftime("%Y-%m-%d %H:%M:%S")
            
            # Handle timestamp (seconds since epoch)
            if date_str.isdigit() and len(date_str) >= 10:
                timestamp = int(date_str)
                dt = datetime.fromtimestamp(timestamp)
                return dt.strftime("%Y-%m-%d %H:%M:%S")
            
            # Handle ISO format
            if 'T' in date_str:
                dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return dt.strftime("%Y-%m-%d %H:%M:%S")
            
            # Return as-is if we can't parse it
            return date_str
            
        except Exception:
            return date_str
    
    def _format_timestamp(self, timestamp: str) -> str:
        """Format timestamp for conversation display"""
        try:
            formatted = self._format_date(timestamp)
            # Extract just time if it's today, or date + time if older
            dt = datetime.fromisoformat(formatted) if 'T' not in formatted else datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
            
            if dt.date() == datetime.now().date():
                return dt.strftime("%H:%M")
            else:
                return dt.strftime("%Y-%m-%d %H:%M")
                
        except Exception:
            return timestamp
    
    def _clean_content(self, content: str) -> str:
        """Clean email content for better readability"""
        if not content:
            return ""
        
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)
        
        # Remove HTML tags if present
        content = re.sub(r'<[^>]+>', '', content)
        
        # Decode HTML entities
        content = content.replace('&amp;', '&')
        content = content.replace('&lt;', '<')
        content = content.replace('&gt;', '>')
        content = content.replace('&quot;', '"')
        content = content.replace('&#39;', "'")
        
        return content.strip()
    
    def _has_conversation_format(self, content: str) -> bool:
        """Check if content already has conversation format with timestamps"""
        # Look for patterns like [timestamp] sender: or [time] name:
        pattern = r'\[.*?\].*?:'
        return bool(re.search(pattern, content))
    
    def _infer_category(self, email_data: Dict[str, Any]) -> str:
        """Infer email category based on content and metadata"""
        subject = email_data.get('subject', '').lower()
        sender_email = email_data.get('sender', {}).get('email', '').lower()
        body = email_data.get('body', '').lower()
        
        # Internal email detection
        internal_domains = ['repsmate.com', 'company.com']  # Add your domains
        if any(domain in sender_email for domain in internal_domains):
            return 'Internal'
        
        # Business category keywords
        business_keywords = ['contract', 'proposal', 'invoice', 'payment', 'business', 'deal', 'partnership']
        if any(keyword in subject or keyword in body for keyword in business_keywords):
            return 'Business'
        
        # Meeting category
        meeting_keywords = ['meeting', 'call', 'schedule', 'appointment', 'calendar']
        if any(keyword in subject or keyword in body for keyword in meeting_keywords):
            return 'Meeting'
        
        # Client category
        client_keywords = ['client', 'customer', 'support', 'issue', 'help']
        if any(keyword in subject or keyword in body for keyword in client_keywords):
            return 'Client'
        
        return 'General'


# Global instance
email_formatter = EmailDocumentFormatter()